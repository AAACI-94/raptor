"""Generate DOCX export from production output."""

from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(payload: dict, title: str) -> BytesIO:
    """Convert production payload to a Word document."""
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document_data = payload.get("document", {})

    # Abstract
    if document_data.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(document_data["abstract"])

    # Sections
    for section in document_data.get("sections", []):
        level = min(section.get("level", 1), 4)
        doc.add_heading(section.get("heading", "Untitled"), level=level)

        content = section.get("content", "")
        # Split content into paragraphs
        for para_text in content.split("\n\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

    # References
    references = document_data.get("references", [])
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            doc.add_paragraph(ref.get("formatted", ""), style="List Number")

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
