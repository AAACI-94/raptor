"""Tests for export endpoints and markdown/docx generators."""

import uuid

from app.models.envelope import ArtifactEnvelope, ArtifactMetadata
from app.services import artifact_service
from app.services.export.markdown_export import generate_markdown
from app.services.export.docx_export import generate_docx


def _create_project(client) -> str:
    """Helper to create a project and return its ID."""
    resp = client.post("/api/projects", json={"title": "Export Test Paper"})
    return resp.json()["id"]


def _store_production_artifact(project_id: str, payload: dict) -> str:
    """Store a production_output artifact and return its ID."""
    envelope = ArtifactEnvelope(
        artifact_id=str(uuid.uuid4()),
        artifact_type="production_output",
        source_agent="production_agent",
        project_id=project_id,
        version=1,
        payload=payload,
        metadata=ArtifactMetadata(model="test-model", duration_ms=100),
        status="submitted",
    )
    return artifact_service.store_artifact(envelope)


def _store_figures_artifact(project_id: str) -> str:
    """Store a figures artifact and return its ID."""
    envelope = ArtifactEnvelope(
        artifact_id=str(uuid.uuid4()),
        artifact_type="figures",
        source_agent="visual_architect",
        project_id=project_id,
        version=1,
        payload={
            "figures": [
                {"id": "fig1", "title": "Architecture Diagram", "mermaid": "graph LR; A-->B"},
                {"id": "fig2", "title": "Data Flow", "mermaid": "graph TD; C-->D"},
            ],
            "figure_plan": "Two diagrams planned",
            "cross_references": ["Section 3 references fig1"],
        },
        metadata=ArtifactMetadata(model="test-model", duration_ms=50),
        status="submitted",
    )
    return artifact_service.store_artifact(envelope)


SAMPLE_PAYLOAD = {
    "document": {
        "title": "Test Paper",
        "abstract": "This is a test abstract for the research paper on security topics.",
        "sections": [
            {"heading": "Introduction", "content": "This paper explores security.", "level": 1},
            {"heading": "Methodology", "content": "We used empirical methods.", "level": 1},
            {"heading": "Background", "content": "Prior work exists.", "level": 2},
        ],
        "references": [
            {"formatted": "[1] Smith et al., Security Research, 2024"},
            {"formatted": "[2] Jones, AI Safety, 2023"},
        ],
    }
}


class TestMarkdownExport:
    """Tests for the markdown generator."""

    def test_generate_markdown_has_title(self):
        """Markdown output should contain the document title as H1."""
        md = generate_markdown(SAMPLE_PAYLOAD, "My Research Paper")
        assert md.startswith("# My Research Paper")

    def test_generate_markdown_has_abstract(self):
        """Markdown output should contain the abstract section."""
        md = generate_markdown(SAMPLE_PAYLOAD, "Test")
        assert "## Abstract" in md
        assert "test abstract" in md

    def test_generate_markdown_has_sections(self):
        """Markdown output should include all document sections."""
        md = generate_markdown(SAMPLE_PAYLOAD, "Test")
        assert "Introduction" in md
        assert "Methodology" in md
        assert "Background" in md

    def test_generate_markdown_has_references(self):
        """Markdown output should include references section."""
        md = generate_markdown(SAMPLE_PAYLOAD, "Test")
        assert "## References" in md
        assert "Smith et al." in md
        assert "Jones" in md

    def test_generate_markdown_section_levels(self):
        """Sections should use correct heading levels based on their level field."""
        md = generate_markdown(SAMPLE_PAYLOAD, "Test")
        # level 1 sections get ## (level + 1 hashes)
        assert "## Introduction" in md
        # level 2 sections get ### (level + 1 hashes)
        assert "### Background" in md

    def test_generate_markdown_empty_payload(self):
        """Markdown output should handle empty payload gracefully."""
        md = generate_markdown({}, "Empty Paper")
        assert "# Empty Paper" in md

    def test_generate_markdown_no_abstract(self):
        """Markdown output should skip abstract section when absent."""
        payload = {"document": {"sections": [{"heading": "Intro", "content": "Hello", "level": 1}]}}
        md = generate_markdown(payload, "Test")
        assert "## Abstract" not in md


class TestDocxExport:
    """Tests for the DOCX generator."""

    def test_generate_docx_returns_bytesio(self):
        """generate_docx should return a BytesIO object."""
        from io import BytesIO
        result = generate_docx(SAMPLE_PAYLOAD, "Test Paper")
        assert isinstance(result, BytesIO)

    def test_generate_docx_is_valid(self):
        """generate_docx output should be readable as a DOCX file."""
        from docx import Document
        buf = generate_docx(SAMPLE_PAYLOAD, "Test Paper")
        doc = Document(buf)
        # Should have paragraphs
        assert len(doc.paragraphs) > 0

    def test_generate_docx_contains_title(self):
        """generate_docx output should contain the title."""
        from docx import Document
        buf = generate_docx(SAMPLE_PAYLOAD, "My Great Paper")
        doc = Document(buf)
        # Title is typically the first paragraph (heading 0)
        texts = [p.text for p in doc.paragraphs]
        assert any("My Great Paper" in t for t in texts)

    def test_generate_docx_empty_payload(self):
        """generate_docx should handle empty payload without crashing."""
        from io import BytesIO
        result = generate_docx({}, "Empty")
        assert isinstance(result, BytesIO)


class TestExportEndpoints:
    """API-level tests for export endpoints."""

    def test_preview_404_no_artifacts(self, client):
        """GET /api/projects/{id}/preview should return 404 when no artifacts exist."""
        project_id = _create_project(client)
        resp = client.get(f"/api/projects/{project_id}/preview")
        assert resp.status_code == 404

    def test_preview_with_production_artifact(self, client):
        """GET /api/projects/{id}/preview should return markdown when production artifact exists."""
        project_id = _create_project(client)
        _store_production_artifact(project_id, SAMPLE_PAYLOAD)

        resp = client.get(f"/api/projects/{project_id}/preview")
        assert resp.status_code == 200
        data = resp.json()
        assert "markdown" in data
        assert data["source"] == "production"
        assert data["word_count"] > 0

    def test_figures_404_no_figures(self, client):
        """GET /api/projects/{id}/figures should return 404 when no figures artifact exists."""
        project_id = _create_project(client)
        resp = client.get(f"/api/projects/{project_id}/figures")
        assert resp.status_code == 404

    def test_figures_with_artifact(self, client):
        """GET /api/projects/{id}/figures should return figure data when artifact exists."""
        project_id = _create_project(client)
        _store_figures_artifact(project_id)

        resp = client.get(f"/api/projects/{project_id}/figures")
        assert resp.status_code == 200
        data = resp.json()
        assert data["total_figures"] == 2
        assert len(data["figures"]) == 2

    def test_export_json_format(self, client):
        """GET /api/projects/{id}/export/json should return the raw payload."""
        project_id = _create_project(client)
        _store_production_artifact(project_id, SAMPLE_PAYLOAD)

        resp = client.get(f"/api/projects/{project_id}/export/json")
        assert resp.status_code == 200
        data = resp.json()
        assert "document" in data

    def test_export_404_no_production(self, client):
        """GET /api/projects/{id}/export/markdown should return 404 without production artifact."""
        project_id = _create_project(client)
        resp = client.get(f"/api/projects/{project_id}/export/markdown")
        assert resp.status_code == 404

    def test_export_unsupported_format(self, client):
        """GET /api/projects/{id}/export/pdf should return 400 for unsupported format."""
        project_id = _create_project(client)
        _store_production_artifact(project_id, SAMPLE_PAYLOAD)

        resp = client.get(f"/api/projects/{project_id}/export/pdf")
        assert resp.status_code == 400

    def test_checklist_404_no_production(self, client):
        """GET /api/projects/{id}/export/checklist should return 404 without production artifact."""
        project_id = _create_project(client)
        resp = client.get(f"/api/projects/{project_id}/export/checklist")
        assert resp.status_code == 404

    def test_export_json_includes_submission_data(self, client):
        """GET /api/projects/{id}/export/json should include submission_checklist in payload."""
        project_id = _create_project(client)
        payload = {**SAMPLE_PAYLOAD, "submission_checklist": ["Check 1", "Check 2"], "formatting_notes": "Use APA"}
        _store_production_artifact(project_id, payload)

        resp = client.get(f"/api/projects/{project_id}/export/json")
        assert resp.status_code == 200
        data = resp.json()
        assert "submission_checklist" in data
        assert len(data["submission_checklist"]) == 2
